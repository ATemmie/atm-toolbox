"""QA: 检查Word文档内容"""
from docx import Document
import os

doc_path = r"\\192.168.0.127\Shared\atm-toolbox\澳门旅行\澳门家庭旅行执行表.docx"
doc = Document(doc_path)

print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print(f"图片数: {sum(1 for rel in doc.part.rels.values() if 'image' in rel.reltype)}")
print()

# 列出所有标题
print("=== 文档结构 ===")
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        level = p.style.name.replace('Heading ', '')
        indent = "  " * int(level) if level.isdigit() else ""
        print(f"{indent}[H{level}] {p.text[:60]}")

# 检查表格内容
print("\n=== 表格概览 ===")
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.cell(0,0).text[:30] if rows > 0 else ""
    print(f"  表格{i+1}: {rows}行×{cols}列 | 首单元格: '{first_cell}'")
