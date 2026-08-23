"""澳门旅行文档内容 - 美化版 v2"""
import os, sys
sys.path.insert(0, os.path.dirname(__file__))

from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from config import DAY1_STOPS, DAY2_STOPS, WEATHER_PLANS, TRANSPORT, TOP10
from doc_builder import (Theme, create_styled_table, add_colored_paragraph,
                         add_divider, add_image_if_exists, set_paragraph_spacing)

IMG_MAP = {
    "a_ma_temple": "妈阁庙", "macau_tower": "澳门旅游塔",
    "senado_square": "议事亭前地", "st_pauls": "大三巴",
    "st_dominic": "玫瑰圣母堂", "monte_fort": "大炮台",
    "nam_van_lake": "南湾湖", "rua_cunha": "官也街",
    "carmel_church": "嘉模圣母堂", "taipa_houses": "龙环葡韵",
    "wynn_palace": "永利皇宫", "parisian": "巴黎人",
    "londoner": "伦敦人", "galaxy": "银河", "venetian": "威尼斯人",
}

# ========== 封面 ==========
def build_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🇲🇴")
    run.font.size = Pt(48)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("澳门家庭旅行攻略")
    run2.font.size = Pt(32)
    run2.bold = True
    run2.font.color.rgb = Theme.PRIMARY
    
    add_divider(doc, "0078D4", 2)
    
    # 副标题
    for text in ["6人家庭 · 2天1夜 · 威尼斯人出发", "2026年8月23日 - 25日"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = Pt(14)
        r.font.color.rgb = Theme.MUTED
    
    doc.add_paragraph()
    
    # 信息卡片
    info_items = [
        ("👨‍👩‍👧‍👦", "3大人 + 2老人 + 1小孩"),
        ("🏨", "住宿：威尼斯人"),
        ("✈️", "23日晚抵达 · 24/25日游玩"),
        ("🚕", "交通：的士为主 + 公交为辅 + 步行"),
    ]
    for icon, text in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{icon}  {text}")
        r.font.size = Pt(11)
        r.font.color.rgb = Theme.TEXT
    
    doc.add_page_break()

# ========== 每日主题概览 ==========
def build_day_type_table(doc):
    doc.add_heading("📅 每日主题概览", level=2)
    
    headers = ["时段", "24日（澳门半岛）", "25日（氹仔+路氹）"]
    rows = [
        ["🌅 上午", "古迹 + 高空景观", "葡式生活 + 风景"],
        ["☀️ 下午", "老城 + 人文", "度假酒店 + 缆车"],
        ["🌙 晚上", "酒店夜景", "路氹夜景"],
    ]
    create_styled_table(doc, headers, rows, "6C5CE7")
    doc.add_paragraph()

# ========== 时间轴表格 ==========
def build_timeline_table(doc, stops, day_label, color_hex):
    doc.add_heading(day_label, level=2)
    doc.add_paragraph()
    
    headers = ["#", "时间", "地点", "说明"]
    rows = []
    for idx, s in enumerate(stops):
        rows.append([str(idx+1), s["time"], s["name"], s["note"]])
    
    create_styled_table(doc, headers, rows, color_hex)
    doc.add_paragraph()

# ========== 景点详情卡片 ==========
def build_attraction_cards(doc, stops):
    doc.add_heading("🏛️ 景点详情", level=3)
    
    for s in stops:
        if s["type"] == "food":
            continue
        
        # 景点标题行
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="0078D4"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
        
        r1 = p.add_run(f"  {s['name']}")
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = Theme.PRIMARY
        r2 = p.add_run(f"    ⏰ {s['time']}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = Theme.MUTED
        
        set_paragraph_spacing(p, 12, 4)
        
        # 图片
        for key, name in IMG_MAP.items():
            if name in s["name"] or s["name"] in name:
                img_path = os.path.join("images", f"{key}.jpg")
                add_image_if_exists(doc, img_path, width=Inches(4.2))
                break
        
        # 备注
        p3 = doc.add_paragraph()
        r3 = p3.add_run(f"💡 {s['note']}")
        r3.font.size = Pt(9)
        r3.font.color.rgb = Theme.MUTED
        r3.italic = True
        set_paragraph_spacing(p3, 0, 8)

# ========== 完整天行程 ==========
def build_day_section(doc, stops, day_label, color_hex, img_keys_map, map_path, plan_key):
    build_timeline_table(doc, stops, day_label, color_hex)
    
    # 路线地图
    doc.add_heading("📍 路线地图", level=3)
    if map_path and os.path.exists(map_path):
        add_image_if_exists(doc, map_path, width=Inches(5.8))
    doc.add_paragraph()
    
    # 景点详情
    build_attraction_cards(doc, stops)
    
    doc.add_page_break()

# ========== 天气备选 ==========
def build_weather_section(doc):
    doc.add_heading("🌧️ 天气备选方案", level=2)
    add_colored_paragraph(doc, "以下为关键时段的三套方案，随时切换：", 
                         font_size=10, color=Theme.MUTED)
    
    headers = ["时段", "☀️ 晴天", "🌧️ 下雨", "⛈️ 暴雨"]
    rows = []
    labels = {
        "day1_afternoon": "24日下午",
        "day2_morning": "25日上午", 
        "day2_afternoon": "25日下午"
    }
    for key, label in labels.items():
        plans = WEATHER_PLANS[key]
        rows.append([label, plans["sunny"], plans["rain"], plans.get("storm", "-")])
    
    create_styled_table(doc, headers, rows, "34495E")
    doc.add_paragraph()

# ========== 交通指南 ==========
def build_transport_section(doc):
    doc.add_heading("🚕 交通指南", level=2)
    add_colored_paragraph(doc, "6人家庭出行，交通优先级排序：", 
                         font_size=10, bold=True)
    
    medals = ["🥇", "🥈", "🥉", "4️⃣", "5️⃣"]
    headers = ["优先级", "方式", "说明"]
    rows = [[medals[i], m, d] for i, (m, d) in enumerate(TRANSPORT["priority"])]
    
    create_styled_table(doc, headers, rows, "2D3436")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("⚠️ 24/25号是周一/周二，不依赖21AT/26AT周末特别线路")
    r.font.size = Pt(10)
    r.font.color.rgb = Theme.DANGER
    r.bold = True
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run("🏨 威尼斯人提供机场/码头/关闸免费穿梭巴士")
    r2.font.size = Pt(10)
    r2.font.color.rgb = Theme.SUCCESS
    
    doc.add_paragraph()

# ========== TOP 10 ==========
def build_top10_section(doc):
    doc.add_heading("⭐ 必去景点 TOP 10", level=2)
    add_colored_paragraph(doc, "如天气极差需砍行程，按此优先级保留：", 
                         font_size=10, color=Theme.MUTED)
    
    headers = ["优先级", "景点"]
    rows = [[stars, name] for stars, name in TOP10]
    
    create_styled_table(doc, headers, rows, "E17055")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("💡 大炮台、南湾湖、SkyCab、湿地属于「天气好就赚到」类型，不强求")
    r.font.size = Pt(9)
    r.font.color.rgb = Theme.MUTED
    r.italic = True
