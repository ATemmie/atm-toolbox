"""Word文档构建器 - 美化版 v2
设计原则：配色方案 + 页眉页脚 + 专业表格 + 卡片式布局
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ========== 配色方案 ==========
class Theme:
    """旅行文档配色方案"""
    PRIMARY = RGBColor(0x1A, 0x1A, 0x2E)      # 深蓝（标题）
    SECONDARY = RGBColor(0x00, 0x78, 0xD4)     # 微软蓝（链接/强调）
    ACCENT = RGBColor(0xE1, 0x70, 0x55)        # 暖橙（Day1/重点）
    ACCENT2 = RGBColor(0x6C, 0x5C, 0xE7)       # 紫色（Day2）
    TEXT = RGBColor(0x33, 0x33, 0x33)           # 正文灰
    MUTED = RGBColor(0x88, 0x88, 0x88)         # 次要文字
    LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)      # 浅灰背景
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    SUCCESS = RGBColor(0x00, 0xB8, 0x94)       # 绿色（晴天）
    WARNING = RGBColor(0xFD, 0xCB, 0x6E)       # 黄色（注意）
    DANGER = RGBColor(0xE1, 0x70, 0x55)        # 红色（警告）
    HEADER_BG = "1A1A2E"                        # 表头背景
    ROW_ALT = "F8F9FA"                          # 交替行背景

# ========== 样式工具 ==========
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=40, start=80, bottom=40, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_paragraph_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line

def add_colored_paragraph(doc, text, font_size=11, bold=False, color=None, 
                          align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    set_paragraph_spacing(p, space_before, space_after)
    return p

def add_divider(doc, color="D0D0D0", width=1):
    """添加水平分割线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width*4}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    set_paragraph_spacing(p, 6, 6)

def add_image_if_exists(doc, img_path, width=Inches(5.5)):
    """添加图片（带验证）"""
    if not img_path or not os.path.exists(img_path):
        return False
    if os.path.getsize(img_path) < 5000:
        return False
    try:
        from PIL import Image as PILImage
        img = PILImage.open(img_path)
        img.verify()
    except Exception:
        return False
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        set_paragraph_spacing(p, 6, 6)
        return True
    except Exception:
        return False

# ========== 表格构建 ==========
def create_styled_table(doc, headers, rows, header_color=Theme.HEADER_BG):
    """创建专业样式表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = Theme.WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 数据行
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        for j, v in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(9)
            r.font.color.rgb = Theme.TEXT
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # 交替行背景
        if idx % 2 == 0:
            for c in row.cells:
                set_cell_shading(c, Theme.ROW_ALT)
    
    return table

# ========== 文档初始化 ==========
def create_document():
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(10.5)
    font.color.rgb = Theme.TEXT
    # 设置东亚字体
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Microsoft YaHei"/>')
        rpr.append(rfonts)
    else:
        rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    
    # 设置标题样式
    for level, (size, color) in {
        1: (22, Theme.PRIMARY),
        2: (16, Theme.PRIMARY),
        3: (13, Theme.SECONDARY),
    }.items():
        hs = doc.styles[f"Heading {level}"]
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.font.bold = True
        hs_rpr = hs.element.get_or_add_rPr()
        hs_rf = hs_rpr.find(qn("w:rFonts"))
        if hs_rf is None:
            hs_rf = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Microsoft YaHei"/>')
            hs_rpr.append(hs_rf)
        else:
            hs_rf.set(qn("w:eastAsia"), "Microsoft YaHei")
    
    return doc

def add_page_header_footer(doc, left_text="澳门家庭旅行攻略", right_text=""):
    """添加页眉页脚"""
    section = doc.sections[0]
    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    run = hp.add_run(left_text)
    run.font.size = Pt(8)
    run.font.color.rgb = Theme.MUTED
    run.font.name = "Microsoft YaHei"
    # 页眉底线
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="D0D0D0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    # 页脚 - 页码
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

def load_images():
    img_dir = os.path.join(os.path.dirname(__file__), "images")
    idx_path = os.path.join(img_dir, "_index.json")
    if os.path.exists(idx_path):
        import json
        with open(idx_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}
