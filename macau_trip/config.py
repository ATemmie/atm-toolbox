"""澳门家庭旅行数据模型 — 2大2老1小 + ATemmie = 6人"""
import os, time, re, urllib.parse, subprocess
from pathlib import Path

PROXY = "http://127.0.0.1:7890"
IMG_DIR = Path(__file__).parent / "images"
MAP_DIR = Path(__file__).parent / "maps"
OUTPUT_DIR = Path(r"\\192.168.0.127\Shared\atm-toolbox\澳门旅行")
FINAL_DOC = OUTPUT_DIR / "澳门家庭旅行执行表.docx"

IMG_DIR.mkdir(exist_ok=True)
MAP_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ── 景点数据 ──────────────────────────────────────────────────
# 每个stop: name, en, lat, lon, time, duration_min, note, type, img_query
DAY1_STOPS = [
    {"name":"威尼斯人","en":"The Venetian Macao","lat":22.1436,"lon":113.5578,
     "time":"22:00","duration_min":60,"note":"入住+酒店内晚餐","type":"hotel",
     "img_query":"威尼斯人酒店 Venetian Macao hotel"},
    {"name":"妈阁庙","en":"A-Ma Temple","lat":22.1895,"lon":113.5320,
     "time":"10:20","duration_min":40,"note":"免费 / 08:00-18:00","type":"heritage",
     "img_query":"妈阁庙 A-Ma Temple Macau"},
    {"name":"澳门旅游塔","en":"Macau Tower","lat":22.1793,"lon":113.5431,
     "time":"11:30","duration_min":90,"note":"只买观光层","type":"attraction",
     "img_query":"澳门旅游塔 Macau Tower"},
    {"name":"南湾湖","en":"Nam Van Lake","lat":22.1888,"lon":113.5425,
     "time":"14:15","duration_min":30,"note":"免费看城市湖景","type":"scenic",
     "img_query":"南湾湖 Nam Van Lake Macau skyline"},
    {"name":"议事亭前地","en":"Senado Square","lat":22.1975,"lon":113.5438,
     "time":"15:00","duration_min":40,"note":"世遗核心 / 步行进入老城","type":"heritage",
     "img_query":"议事亭前地 Senado Square Macau"},
    {"name":"玫瑰圣母堂","en":"St. Dominic's Church","lat":22.1995,"lon":113.5434,
     "time":"15:40","duration_min":20,"note":"免费","type":"heritage",
     "img_query":"玫瑰圣母堂 St Dominic Church Macau"},
    {"name":"大三巴","en":"Ruins of St. Paul's","lat":22.1987,"lon":113.5405,
     "time":"16:20","duration_min":40,"note":"免费 / 含哪咤庙+恋爱巷","type":"heritage",
     "img_query":"大三巴牌坊 Ruins of St Pauls Macau"},
    {"name":"大炮台","en":"Monte Fort","lat":22.1985,"lon":113.5438,
     "time":"17:00","duration_min":40,"note":"☀️晴天去 / 可看城市全景","type":"heritage",
     "img_query":"大炮台 Monte Fort Macau"},
    {"name":"晚餐","en":"Dinner near Senado","lat":22.1970,"lon":113.5435,
     "time":"18:00","duration_min":90,"note":"大三巴/新马路附近","type":"food",
     "img_query":"澳门美食 Macau food egg tart pork chop bun"},
    {"name":"巴黎人","en":"The Parisian Macao","lat":22.1478,"lon":113.5606,
     "time":"20:30","duration_min":30,"note":"巴黎铁塔夜景","type":"hotel",
     "img_query":"巴黎人酒店 Parisian Macao Eiffel Tower"},
    {"name":"伦敦人","en":"The Londoner Macao","lat":22.1474,"lon":113.5590,
     "time":"21:00","duration_min":30,"note":"大本钟+英伦建筑夜景","type":"hotel",
     "img_query":"伦敦人酒店 Londoner Macao"},
]

DAY2_STOPS = [
    {"name":"威尼斯人","en":"The Venetian Macao","lat":22.1436,"lon":113.5578,
     "time":"10:15","duration_min":0,"note":"出发","type":"hotel",
     "img_query":"威尼斯人酒店 Venetian Macao canal"},
    {"name":"官也街","en":"Rua do Cunha","lat":22.1523,"lon":113.5514,
     "time":"10:30","duration_min":60,"note":"小吃+手信 / 不安排正餐","type":"food",
     "img_query":"官也街 Rua do Cunha Taipa Macau"},
    {"name":"嘉模圣母堂","en":"Chapel of Our Lady of Carmel","lat":22.1537,"lon":113.5498,
     "time":"11:30","duration_min":20,"note":"免费","type":"heritage",
     "img_query":"嘉模圣母堂 Chapel Our Lady Carmel Taipa"},
    {"name":"龙环葡韵","en":"Taipa Houses-Museum","lat":22.1552,"lon":113.5508,
     "time":"11:50","duration_min":60,"note":"☀️晴天去 / 澳门八景之一","type":"scenic",
     "img_query":"龙环葡韵 Taipa Houses Macau Portuguese"},
    {"name":"午餐","en":"Lunch Taipa","lat":22.1530,"lon":113.5510,
     "time":"13:00","duration_min":60,"note":"提前预约坐得舒服的餐厅","type":"food",
     "img_query":"澳门葡国菜 Macanese food restaurant"},
    {"name":"威尼斯人休息","en":"Rest at Venetian","lat":22.1436,"lon":113.5578,
     "time":"14:00","duration_min":90,"note":"老人午睡 / 小孩玩手机","type":"hotel",
     "img_query":"威尼斯人酒店套房 Venetian Macao suite"},
    {"name":"永利皇宫","en":"Wynn Palace","lat":22.1499,"lon":113.5620,
     "time":"16:00","duration_min":60,"note":"湖边+花艺+水景","type":"hotel",
     "img_query":"永利皇宫 Wynn Palace Macau performance lake"},
    {"name":"SkyCab缆车","en":"SkyCab Wynn Palace","lat":22.1499,"lon":113.5620,
     "time":"17:00","duration_min":30,"note":"☀️晴天坐 / ⛈️雷雨取消","type":"attraction",
     "img_query":"永利皇宫缆车 SkyCab Wynn Palace Macau"},
    {"name":"银河","en":"Galaxy Macau","lat":22.1523,"lon":113.5623,
     "time":"17:30","duration_min":60,"note":"看建筑+甜品+休息","type":"hotel",
     "img_query":"银河度假城 Galaxy Macau resort"},
    {"name":"巴黎人","en":"The Parisian Macao","lat":22.1478,"lon":113.5606,
     "time":"18:30","duration_min":30,"note":"巴黎铁塔拍照","type":"hotel",
     "img_query":"巴黎人巴黎铁塔 Parisian Macao Eiffel Tower night"},
    {"name":"伦敦人","en":"The Londoner Macao","lat":22.1474,"lon":113.5590,
     "time":"19:00","duration_min":30,"note":"大本钟+英伦建筑","type":"hotel",
     "img_query":"伦敦人大本钟 Londoner Macao Big Ben"},
    {"name":"晚餐","en":"Dinner Cotai","lat":22.1480,"lon":113.5590,
     "time":"20:00","duration_min":60,"note":"伦敦人/威尼斯人内吃","type":"food",
     "img_query":"澳门威尼斯人美食 Venetian Macao dining"},
]

TYPE_COLORS = {
    "hotel":"#6C5CE7","heritage":"#E17055","attraction":"#00B894",
    "scenic":"#0984E3","food":"#FDCB6E",
}

# ── 图片下载 ──────────────────────────────────────────────────
def download_image(query: str, filename: str) -> str | None:
    """从Bing搜索下载景点图片，返回本地路径或None"""
    filepath = IMG_DIR / filename
    if filepath.exists() and filepath.stat().st_size > 5000:
        print(f"  [cached] {filename}")
        return str(filepath)

    encoded = urllib.parse.quote(query)
    search_url = f"https://cn.bing.com/images/search?q={encoded}&first=1&count=5"
    try:
        result = subprocess.run(
            ["curl", "-sL", "--proxy", PROXY,
             "-H", "User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
             "-H", "Referer: https://www.bing.com/",
             "--connect-timeout", "10", "--max-time", "15",
             search_url],
            capture_output=True, text=True, timeout=20
        )
        html = result.stdout
        patterns = [
            r'"murl":"(https?://[^"]+)"',
            r'"purl":"(https?://[^"]+)"',
            r'src2="(https?://[^"]+\.(?:jpg|jpeg|png))"',
            r'data-src="(https?://[^"]+\.(?:jpg|jpeg|png))"',
        ]
        img_url = None
        for pat in patterns:
            urls = re.findall(pat, html)
            urls = [re.sub(r'[&].*$', '', u) for u in urls]
            urls = [u for u in urls if 'bing.com' not in u and 'microsoft' not in u.lower()]
            if urls:
                img_url = urls[0]
                break

        if not img_url:
            all_urls = re.findall(r'https?://[^"\'\\s<>]+', html)
            img_url = next((u for u in all_urls
                           if any(ext in u.lower() for ext in ['.jpg','.jpeg','.png'])
                           and 'bing.com' not in u and 'microsoft' not in u.lower()), None)

        if not img_url:
            print(f"  [no url] {query}")
            return None

        dl = subprocess.run(
            ["curl", "-sL", "-o", str(filepath),
             "-H", "User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
             "-H", "Referer: https://www.bing.com/",
             "--proxy", PROXY, "--connect-timeout", "10", "--max-time", "20",
             img_url],
            capture_output=True, timeout=25
        )
        # 验证
        from PIL import Image as PILImage
        if filepath.exists() and filepath.stat().st_size > 5000:
            try:
                img = PILImage.open(filepath)
                img.verify()
                print(f"  [ok] {filename} ({filepath.stat().st_size//1024}KB)")
                return str(filepath)
            except Exception:
                filepath.unlink(missing_ok=True)
                print(f"  [invalid] {query}")
                return None
        else:
            filepath.unlink(missing_ok=True)
            print(f"  [too small] {query}")
            return None
    except Exception as e:
        print(f"  [error] {query}: {e}")
        return None

def download_all_images():
    """下载所有景点图片"""
    all_stops = DAY1_STOPS + DAY2_STOPS
    seen = set()
    results = {}
    for stop in all_stops:
        key = stop["en"]
        if key in seen:
            results[key] = results.get(key)
            continue
        seen.add(key)
        safe_name = re.sub(r'[^\w]', '_', stop["en"])[:40] + ".jpg"
        print(f"Downloading: {stop['name']} ({stop['en']})...")
        path = download_image(stop["img_query"], safe_name)
        results[key] = path
        time.sleep(1)  # rate limit
    return results

# ── 地图生成 ──────────────────────────────────────────────────
MAP_HTML_TEMPLATE = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css"/>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<style>
body {{ margin:0; padding:0; }}
#map {{ width:1200px; height:750px; }}
.legend {{
    position:absolute; bottom:20px; right:20px; z-index:1000;
    background:rgba(255,255,255,0.95); padding:10px 14px; border-radius:8px;
    font-family:Arial,sans-serif; font-size:12px; box-shadow:0 2px 6px rgba(0,0,0,0.3);
}}
.legend-item {{ display:flex; align-items:center; margin:3px 0; }}
.legend-dot {{ width:12px; height:12px; border-radius:50%; margin-right:6px; border:1px solid #fff; }}
.title-bar {{
    position:absolute; top:10px; left:50%; transform:translateX(-50%); z-index:1000;
    background:rgba(30,30,30,0.85); color:#fff; padding:8px 20px; border-radius:20px;
    font-family:Arial,sans-serif; font-size:16px; font-weight:bold;
    box-shadow:0 2px 8px rgba(0,0,0,0.3);
}}
</style></head><body>
<div class="title-bar">{title}</div>
<div id="map"></div>
<div class="legend">
{legend_items}
</div>
<script>
var map = L.map('map',{{zoomControl:true}}).setView([{center_lat},{center_lon}],{zoom});
L.tileLayer('https://{{s}}.tile.openstreetmap.org/{{z}}/{{x}}/{{y}}.png', {{
    attribution:'© OpenStreetMap', maxZoom:19
}}).addTo(map);

var stops = {stops_json};
var coords = [];
stops.forEach(function(s,i){{
    L.circleMarker([s.lat,s.lon],{{
        radius:14, fillColor:s.color, color:'#fff', weight:2.5, fillOpacity:0.9
    }}).addTo(map).bindTooltip(String(i+1),{{permanent:true, direction:'center', className:'num-label'}});
    coords.push([s.lat,s.lon]);
}});

L.polyline(coords,{{color:'#E17055', weight:4, opacity:0.85, dashArray:'10,6'}}).addTo(map);
map.fitBounds(coords,{{padding:[60,60]}});
</script>
</body></html>"""

def generate_map_html(day_num, stops, title):
    """生成Leaflet地图HTML"""
    legend_items = "\n".join(
        f'<div class="legend-item"><div class="legend-dot" style="background:{TYPE_COLORS[s["type"]]}"></div>{s["name"]}</div>'
        for s in stops
    )
    stops_data = [
        {"lat":s["lat"],"lon":s["lon"],"color":TYPE_COLORS[s["type"]],"name":s["name"]}
        for s in stops
    ]
    import json
    center_lat = sum(s["lat"] for s in stops)/len(stops)
    center_lon = sum(s["lon"] for s in stops)/len(stops)
    html = MAP_HTML_TEMPLATE.format(
        title=title,
        center_lat=center_lat, center_lon=center_lon, zoom=14,
        stops_json=json.dumps(stops_data, ensure_ascii=False),
        legend_items=legend_items
    )
    html_path = MAP_DIR / f"day{day_num}_map.html"
    html_path.write_text(html, encoding="utf-8")
    return str(html_path)

async def screenshot_map(html_path, png_path):
    """Playwright截图地图"""
    from playwright.async_api import async_playwright
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page(viewport={"width":1200,"height":750})
        file_url = "file:///" + os.path.abspath(html_path).replace("\\","/")
        await page.goto(file_url, wait_until="networkidle")
        await page.wait_for_timeout(3000)
        await page.screenshot(path=png_path, full_page=False)
        await browser.close()
    print(f"  Map saved: {png_path}")

# ── Word文档组装 ───────────────────────────────────────────────
def build_document(day1_map_png, day2_map_png, img_dict):
    """组装Word文档"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from PIL import Image as PILImage

    doc = Document()

    # ── 页面设置 A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    def set_cell_shading(cell, color):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color[1:]}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_styled_heading(text, level=1, color=None):
        h = doc.add_heading(text, level=level)
        if color:
            for run in h.runs:
                run.font.color.rgb = color
        return h

    def add_image_safe(path, width_inches=6):
        """安全插入图片"""
        if path and os.path.exists(path) and os.path.getsize(path) > 5000:
            try:
                img = PILImage.open(path)
                img.verify()
                doc.add_picture(path, width=Inches(width_inches))
                return True
            except Exception:
                pass
        return False

    def add_placeholder_card(name, en):
        """文字占位卡"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"📷 {name}\n{en}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(128,128,128)

    # ══════════════════════════════════════════════════════════
    # 封面
    # ══════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("🇲🇴 澳门家庭旅行执行表")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("6人家庭 · 2天1夜 · 威尼斯人出发")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("2026年8月23日-25日\n3大人 + 2老人 + 1小孩\n入住威尼斯人")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 总览
    # ══════════════════════════════════════════════════════════
    add_styled_heading("行程总览", level=1, color=RGBColor(30,64,175))

    overview_data = [
        ["日期","主题","关键词","核心景点"],
        ["8月24日(一)","澳门半岛\n南→北","历史+海景+世遗","澳门塔/大三巴/议事亭"],
        ["8月25日(二)","氹仔+路氹\n东南→西北","葡式建筑+酒店夜景","官也街/龙环葡韵/永利"],
    ]
    table = doc.add_table(rows=len(overview_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(overview_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)
            if i == 0:
                set_cell_shading(cell, "1E40AF")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True
            elif i % 2 == 0:
                set_cell_shading(cell, "EFF6FF")

    doc.add_paragraph()

    # 时段概览
    add_styled_heading("每日时段安排", level=2)
    time_slots = [
        ["时段","24日(半岛)","25日(氹仔)"],
        ["上午 9-12","妈阁庙→澳门塔","官也街→嘉模圣母堂→龙环葡韵"],
        ["下午 12-17","午餐→南湾湖→议事亭→大三巴","午餐→威尼斯人休息"],
        ["晚上 17-22","大炮台→晚餐→巴黎人伦敦人","永利→银河→巴黎人→伦敦人→晚餐"],
    ]
    table = doc.add_table(rows=len(time_slots), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(time_slots):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(cell, "059669")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # DAY 1
    # ══════════════════════════════════════════════════════════
    add_styled_heading("📅 8月24日(周一)：澳门半岛南→北", level=1, color=RGBColor(225,112,85))

    # 路线图
    add_styled_heading("🗺️ 路线地图", level=2)
    if day1_map_png and os.path.exists(day1_map_png):
        doc.add_picture(day1_map_png, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 时间轴
    add_styled_heading("⏰ 时间轴", level=2)
    timeline_data = [["#","时间","地点","时长","备注"]]
    for i, s in enumerate(DAY1_STOPS, 1):
        timeline_data.append([
            str(i), s["time"], s["name"],
            f'{s["duration_min"]}分钟' if s["duration_min"] > 0 else "-",
            s["note"]
        ])
    table = doc.add_table(rows=len(timeline_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(timeline_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, "E17055")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True
            elif i % 2 == 0:
                set_cell_shading(cell, "FFF5F0")

    doc.add_paragraph()

    # 景点详情卡
    add_styled_heading("📍 景点详情", level=2)
    for s in DAY1_STOPS:
        # 名称 + 类型标签
        p = doc.add_paragraph()
        type_emoji = {"hotel":"🏨","heritage":"🏛️","attraction":"🎡","scenic":"🌊","food":"🍽️"}
        run = p.add_run(f'{type_emoji.get(s["type"],"")} {s["name"]}')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30,30,30)

        en_p = doc.add_paragraph()
        run = en_p.add_run(s["en"])
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128,128,128)
        run.font.italic = True

        # 图片
        img_path = img_dict.get(s["en"])
        if img_path:
            add_image_safe(img_path, width_inches=4.5)
        else:
            add_placeholder_card(s["name"], s["en"])

        # 信息
        info_text = f'⏰ {s["time"]}  ·  ⏱️ {s["duration_min"]}分钟\n{s["note"]}'
        info_p = doc.add_paragraph()
        run = info_p.add_run(info_text)
        run.font.size = Pt(10)

        doc.add_paragraph()  # spacing

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # DAY 2
    # ══════════════════════════════════════════════════════════
    add_styled_heading("📅 8月25日(周二)：氹仔+路氹", level=1, color=RGBColor(9,132,227))

    add_styled_heading("🗺️ 路线地图", level=2)
    if day2_map_png and os.path.exists(day2_map_png):
        doc.add_picture(day2_map_png, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_styled_heading("⏰ 时间轴", level=2)
    timeline_data = [["#","时间","地点","时长","备注"]]
    for i, s in enumerate(DAY2_STOPS, 1):
        timeline_data.append([
            str(i), s["time"], s["name"],
            f'{s["duration_min"]}分钟' if s["duration_min"] > 0 else "-",
            s["note"]
        ])
    table = doc.add_table(rows=len(timeline_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(timeline_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, "0984E3")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True
            elif i % 2 == 0:
                set_cell_shading(cell, "EBF5FB")

    doc.add_paragraph()

    add_styled_heading("📍 景点详情", level=2)
    for s in DAY2_STOPS:
        p = doc.add_paragraph()
        type_emoji = {"hotel":"🏨","heritage":"🏛️","attraction":"🎡","scenic":"🌊","food":"🍽️"}
        run = p.add_run(f'{type_emoji.get(s["type"],"")} {s["name"]}')
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30,30,30)

        en_p = doc.add_paragraph()
        run = en_p.add_run(s["en"])
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128,128,128)
        run.font.italic = True

        img_path = img_dict.get(s["en"])
        if img_path:
            add_image_safe(img_path, width_inches=4.5)
        else:
            add_placeholder_card(s["name"], s["en"])

        info_text = f'⏰ {s["time"]}  ·  ⏱️ {s["duration_min"]}分钟\n{s["note"]}'
        info_p = doc.add_paragraph()
        run = info_p.add_run(info_text)
        run.font.size = Pt(10)

        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 备用路线
    # ══════════════════════════════════════════════════════════
    add_styled_heading("🌧️ 天气备用路线", level=1, color=RGBColor(214,48,49))

    backup_data = [
        ["时段","☀️ 晴天","🌧️ 下雨","⛈️ 雷暴"],
        ["24日上午","妈阁庙→澳门塔","妈阁庙→澳门塔(不变)","妈阁庙→澳门塔"],
        ["24日下午","南湾湖→议事亭→大三巴→大炮台","澳门塔→餐厅→议事亭→大三巴→澳门博物馆","餐厅→威尼斯人"],
        ["25日上午","官也街→嘉模圣母堂→龙环葡韵→湿地","官也街→室内商业区→午餐","威尼斯人室内→午餐"],
        ["25日下午","永利→SkyCab→银河→巴黎人→伦敦人","永利室内→银河→巴黎人→伦敦人","威尼斯人室内→巴黎人→伦敦人"],
    ]
    table = doc.add_table(rows=len(backup_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(backup_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if i == 0:
                set_cell_shading(cell, "7F1D1D")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 交通优先级
    # ══════════════════════════════════════════════════════════
    add_styled_heading("🚕 交通优先级", level=1, color=RGBColor(102,51,153))

    transport_data = [
        ["优先级","方式","适用场景","说明"],
        ["🥇 第一","酒店免费接驳","机场→酒店","威尼斯人免费穿梭巴士"],
        ["🥈 第二","6人特别的士","远距离移动","6人一起走"],
        ["🥉 第三","普通的士×2","叫不到6人车时","两辆普通出租车"],
        ["4️⃣","公交","天气好+不赶时间","26A/25/15等"],
        ["5️⃣","步行","景点间很近","议事亭→大三巴段"],
    ]
    table = doc.add_table(rows=len(transport_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(transport_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(cell, "6B21A8")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 核心景点TOP 10
    # ══════════════════════════════════════════════════════════
    add_styled_heading("⭐ 必去景点 TOP 10", level=1, color=RGBColor(217,119,6))

    top10_data = [
        ["排名","景点","评分","天气依赖"],
        ["1","澳门旅游塔","⭐⭐⭐⭐⭐","室内，不受天气影响"],
        ["2","大三巴","⭐⭐⭐⭐⭐","室外，但下雨可快速拍照离开"],
        ["3","议事亭前地","⭐⭐⭐⭐⭐","室外，下雨可快速穿过"],
        ["4","官也街","⭐⭐⭐⭐⭐","半室外，有骑楼遮雨"],
        ["5","龙环葡韵","⭐⭐⭐⭐⭐","☀️晴天强推 / ⛈️雷雨取消"],
        ["6","妈阁庙","⭐⭐⭐⭐","室外，30分钟可完成"],
        ["7","永利皇宫","⭐⭐⭐⭐","室内为主"],
        ["8","巴黎人","⭐⭐⭐⭐","室内外结合"],
        ["9","伦敦人","⭐⭐⭐⭐","室内为主"],
        ["10","银河","⭐⭐⭐⭐","室内为主"],
    ]
    table = doc.add_table(rows=len(top10_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(top10_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(cell, "92400E")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════
    # 预算估算
    # ══════════════════════════════════════════════════════════
    add_styled_heading("💰 预算估算 (6人2天)", level=1, color=RGBColor(5,150,105))

    budget_data = [
        ["项目","单价(澳门币)","人数/次","小计(MOP)","备注"],
        ["澳门塔观光层","~165","6","~990","建议只买观光层"],
        ["大炮台/博物馆","~15","6","~90","如有兴趣"],
        ["SkyCab缆车","~100","6","~600","☀️晴天可选"],
        ["的士(预估6趟)","~30/趟","6趟","~180","6人特别的士为主"],
        ["公交(预估)","~6","10人次","~60","备用"],
        ["午餐×2","~150/人","6×2","~1800","含老人小孩"],
        ["晚餐×2","~200/人","6×2","~2400","含老人小孩"],
        ["小吃/手信","~200/人","6","~1200","官也街+大三巴"],
        ["","","","",""],
        ["合计估算","","","~7,140 MOP","≈ ¥6,400"],
        ["","","","",""],
        ["注：住宿和机场交通已另计","","","",""],
    ]
    table = doc.add_table(rows=len(budget_data), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(budget_data):
        for j, text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if i == 0:
                set_cell_shading(cell, "065F46")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255,255,255)
                        run.font.bold = True
            elif i == len(budget_data) - 3:  # 合计行
                set_cell_shading(cell, "D1FAE5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 实用提醒
    # ══════════════════════════════════════════════════════════
    add_styled_heading("📋 实用提醒", level=1, color=RGBColor(55,65,81))

    tips = [
        "🎫 景点门票：大部分景点免费，澳门塔观光层+SkyCab是主要门票支出",
        "🌧️ 天气应对：8月底澳门仍属雨季，随身带折叠伞",
        "👴 老人关怀：24日午后大三巴段+25日下午安排休息，避免暴走",
        "👶 小孩照顾：官也街小吃可安抚，威尼斯人室内有活动空间",
        "📱 通讯：澳门多数酒店/商场有免费WiFi",
        "💳 支付：澳门大部分地方支持支付宝/微信/信用卡，少量小店需港币现金",
        "🚌 公交注意：24、25日是周一/周二，不能用周末特别线路(21AT/26AT)",
        "🚕 打车：澳门的士起步价MOP19，6人需叫'特别的士'(加价)",
        "⏰ 时差：澳门与北京时间相同(UTC+8)",
        "🔌 插座：澳门用英标三脚插头(与内地不同)，需转换器",
    ]
    for tip in tips:
        p = doc.add_paragraph(tip)
        for run in p.runs:
            run.font.size = Pt(10)

    # ── 保存 ──
    doc.save(str(FINAL_DOC))
    print(f"\n✅ Document saved: {FINAL_DOC}")
    return str(FINAL_DOC)

# ── 主流程 ──────────────────────────────────────────────────
if __name__ == "__main__":
    import asyncio

    print("="*60)
    print("澳门家庭旅行执行表 — 文档生成")
    print("="*60)

    # 1. 下载图片
    print("\n[1/4] 下载景点图片...")
    img_dict = download_all_images()
    ok_count = sum(1 for v in img_dict.values() if v)
    print(f"\n图片下载完成: {ok_count}/{len(img_dict)} 成功")

    # 2. 生成地图HTML
    print("\n[2/4] 生成路线地图...")
    day1_html = generate_map_html(1, DAY1_STOPS, "Day 1: 澳门半岛 南→北")
    day2_html = generate_map_html(2, DAY2_STOPS, "Day 2: 氹仔+路氹")
    print(f"  HTML: {day1_html}")
    print(f"  HTML: {day2_html}")

    # 3. Playwright截图
    print("\n[3/4] 截图地图...")
    day1_png = str(MAP_DIR / "day1_map.png")
    day2_png = str(MAP_DIR / "day2_map.png")
    asyncio.run(screenshot_map(day1_html, day1_png))
    asyncio.run(screenshot_map(day2_html, day2_png))

    # 4. 组装Word文档
    print("\n[4/4] 组装Word文档...")
    output = build_document(day1_png, day2_png, img_dict)

    print("\n" + "="*60)
    print("DONE!")
    print(f"Output: {output}")
    print("="*60)
